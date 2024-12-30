from django.http import JsonResponse, FileResponse
from django.shortcuts import render
from PIL import Image
from io import BytesIO
import os



def index(request):
    return render(request, 'index.html')


def img_to_pdf(request):
    if request.method == 'POST':
        files = request.FILES.getlist('images')

        # Ensure at least one file is uploaded
        if len(files) < 1:
            return JsonResponse({'error': 'Please upload at least one image.'}, status=400)

        images = []
        for file in files:
            try:
                # Open image and ensure it's in RGB mode for PDF compatibility
                img = Image.open(file)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                images.append(img)
            except Exception as e:
                return JsonResponse({'error': f'Error processing file {file.name}: {e}'}, status=400)

        # Create a PDF from the images
        pdf_file = BytesIO()
        first_image = images.pop(0)
        first_image.save(pdf_file, format='PDF', save_all=True, append_images=images)
        pdf_file.seek(0)

        # Generate a file name for the PDF
        first_file_name = os.path.splitext(files[0].name)[0]
        pdf_file_name = f"{first_file_name}_converted.pdf"

        # Return the PDF as a file response
        response = FileResponse(pdf_file, as_attachment=True, filename=pdf_file_name)
        return response

    return render(request, 'img_to_pdf.html')
  


# views.py
import os
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas

@csrf_exempt
def word_to_pdf(request):
    if request.method == "POST" and request.FILES.get("word_files"):
        files = request.FILES.getlist("word_files")
        pdf_buffer = BytesIO()
        canvas_obj = canvas.Canvas(pdf_buffer)

        try:
            for file in files:
                if file.name.endswith(".docx"):
                    doc = Document(file)
                    
                    # Start writing from the top of the page
                    x, y = 50, 800  
                    
                    for paragraph in doc.paragraphs:
                        # Split paragraph into lines if it exceeds the page width
                        lines = canvas_obj.beginText(x, y)
                        lines.setFont("Helvetica", 12)
                        
                        # Handle word wrapping manually for long text
                        max_width = 500  # Adjust based on page width
                        words = paragraph.text.split()
                        line = ""
                        
                        for word in words:
                            if canvas_obj.stringWidth(line + word + " ", "Helvetica", 12) <= max_width:
                                line += word + " "
                            else:
                                lines.textLine(line.strip())
                                line = word + " "
                                y -= 20
                                if y < 50:  # Move to a new page if needed
                                    canvas_obj.drawText(lines)
                                    canvas_obj.showPage()
                                    y = 800
                                    lines = canvas_obj.beginText(x, y)
                        
                        if line:
                            lines.textLine(line.strip())
                            y -= 20

                        # Ensure proper spacing between paragraphs
                        y -= 10
                        if y < 50:  # Move to a new page if needed
                            canvas_obj.drawText(lines)
                            canvas_obj.showPage()
                            y = 800
                            lines = canvas_obj.beginText(x, y)

                        canvas_obj.drawText(lines)
                else:
                    return JsonResponse({"error": "Invalid file format. Only .docx files are allowed."}, status=400)

            canvas_obj.save()
            pdf_buffer.seek(0)

            response = HttpResponse(pdf_buffer, content_type="application/pdf")
            response["Content-Disposition"] = f'attachment; filename="converted_files.pdf"'
            return response
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return render(request, "word_to_pdf.html")

# import os
# from io import BytesIO
# from django.http import HttpResponse, JsonResponse
# from django.shortcuts import render
# from django.views.decorators.csrf import csrf_exempt
# from pptx import Presentation
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter

# @csrf_exempt
# def ppt_to_pdf(request):
    if request.method == "POST" and request.FILES.get("ppt_files"):
        files = request.FILES.getlist("ppt_files")
        pdf_buffer = BytesIO()

        try:
            pdf_canvas = canvas.Canvas(pdf_buffer, pagesize=letter)

            for file in files:
                if file.name.endswith(".pptx"):
                    # Use BytesIO stream for the uploaded file
                    pptx_file = BytesIO(file.read())
                    prs = Presentation(pptx_file)

                    # Iterate through each slide
                    for slide in prs.slides:
                        y_position = 750  # Start at the top of the page
                        for shape in slide.shapes:
                            if shape.has_text_frame:
                                for paragraph in shape.text_frame.paragraphs:
                                    text = paragraph.text
                                    pdf_canvas.drawString(50, y_position, text)
                                    y_position -= 20  # Move down for the next line
                                    if y_position < 50:  # Start a new page if space runs out
                                        pdf_canvas.showPage()
                                        y_position = 750
                        pdf_canvas.showPage()  # Move to the next page for the next slide
                else:
                    return JsonResponse({"error": "Invalid file format. Only .pptx files are allowed."}, status=400)

            pdf_canvas.save()
            pdf_buffer.seek(0)

            response = HttpResponse(pdf_buffer, content_type="application/pdf")
            response["Content-Disposition"] = f'attachment; filename="converted_files.pdf"'
            return response
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)

    return render(request, 'ppt_to_pdf.html')



# import os
# from io import BytesIO
# from django.http import HttpResponse, JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from pdf2docx import Converter
# from tempfile import NamedTemporaryFile

# @csrf_exempt
# def pdf_to_word(request):
#     if request.method == "POST" and request.FILES.get("pdf_file"):
#         pdf_file = request.FILES["pdf_file"]

#         try:
#             # Create a temporary file for the uploaded PDF
#             with NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
#                 temp_pdf.write(pdf_file.read())
#                 temp_pdf_path = temp_pdf.name

#             # Create a temporary file for the Word document
#             with NamedTemporaryFile(delete=False, suffix=".docx") as temp_word:
#                 temp_word_path = temp_word.name

#             # Convert PDF to Word
#             converter = Converter(temp_pdf_path)
#             converter.convert(temp_word_path)  # Convert to Word
#             converter.close()

#             # Read the Word file and send it in response
#             with open(temp_word_path, "rb") as word_file:
#                 response = HttpResponse(word_file.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
#                 response["Content-Disposition"] = f'attachment; filename="converted_document.docx"'

#             # Clean up temporary files
#             os.remove(temp_pdf_path)
#             os.remove(temp_word_path)

#             return response
#         except Exception as e:
#             return JsonResponse({"error": str(e)}, status=500)

#     return render(request,'pdf_to_word.html')

   
