from django.shortcuts import render
from django.http import HttpResponse
from PyPDF2 import PdfMerger


def index(request):
    return render(request, 'index.html')

def merge_pdfs(request):
    if request.method == "POST":
        files = request.FILES.getlist('pdfs')
        if not files:
            return HttpResponse("No files uploaded.", status=400)

        merger = PdfMerger()
        first_file_name = files[0].name.split('.')[0]  # Extract the name of the first file (without extension)

        for file in files:
            merger.append(file)

        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{first_file_name}_merged.pdf"'
        merger.write(response)
        merger.close()

        return response

    return render(request, 'merge/merge_pdfs.html')

from django.http import JsonResponse, FileResponse
from django.shortcuts import render
from docx import Document
from io import BytesIO


def merge_words(request):
    if request.method == 'POST':
        files = request.FILES.getlist('words')

        # Ensure at least two files are uploaded
        if len(files) < 2:
            return JsonResponse({'error': 'Please upload at least two Word files.'}, status=400)

        # Create a new Word document to merge into
        merged_document = Document()

        for file in files:
            if file.name.endswith('.docx'):
                doc = Document(file)
                for paragraph in doc.paragraphs:
                    merged_document.add_paragraph(paragraph.text)
            else:
                return JsonResponse({'error': f'{file.name} is not a valid Word document.'}, status=400)

        # Save merged document to a BytesIO object
        merged_file = BytesIO()
        merged_document.save(merged_file)
        merged_file.seek(0)

        # Generate merged file name
        first_file_name = files[0].name.split('.')[0]
        merged_file_name = f"{first_file_name}_merged.docx"

        # Return the merged file as a response
        response = FileResponse(merged_file, as_attachment=True, filename=merged_file_name)
        return response

    return render(request, 'merge/merge_words.html')
    
